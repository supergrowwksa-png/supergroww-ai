"""
Candidate Management Endpoints with Production Logging
======================================================

Handles all candidate-related operations with comprehensive logging:
- Creating candidates
- Uploading/parsing resumes
- Listing candidates
- Retrieving candidate details

LOGGING FEATURES IMPLEMENTED:
1. Request/response tracking with context
2. File upload processing metrics
3. Resume parsing performance
4. Embedding generation tracking
5. Validation error logging
6. Database operation monitoring
"""

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from sqlalchemy.orm import Session
from typing import List, Optional
import PyPDF2
import docx
import io
import uuid

from backend.database import (
    get_session, Candidate, Job,
    create_candidate_with_resume
)
from backend.app.api.schemas.models import (
    CreateCandidateRequest, CandidateResponse, ErrorResponse
)

# PROPER LOGGING SETUP (Following logger.py patterns)

from backend.logger import (
    get_logger,
    ContextLogger,
    PerformanceLogger,
    log_exception
)

# Create component logger
logger = get_logger("api.candidates")

router = APIRouter(prefix="/api/candidates", tags=["Candidates"])

# CREATE CANDIDATE

@router.post("/", response_model=CandidateResponse, status_code=status.HTTP_201_CREATED)
async def create_candidate(
    request: CreateCandidateRequest,
    db: Session = Depends(get_session)
):
    """
    Create a new candidate with comprehensive logging.
    
    Flow:
    1. Validate job exists
    2. Check if candidate already exists (by email)
    3. Create candidate record
    4. If resume text provided, generate embeddings
    """
    
    # Generate request ID for tracing
    request_id = f"req-{uuid.uuid4().hex[:12]}"
    
    ctx_logger = ContextLogger(
        logger=logger,
        request_id=request_id,
        candidate_email=request.email,
        job_id=request.job_id
    )
    
    ctx_logger.info(
        "Creating new candidate",
        extra={
            "candidate_name": request.name,
            "has_resume": bool(request.resume_text),
            "resume_length": len(request.resume_text) if request.resume_text else 0,
            "source": request.source
        }
    )
    
    try:
        # 1. Validate job exists
        ctx_logger.debug("Validating job exists")
        
        with PerformanceLogger(logger, "validate_job", request_id=request_id):
            job = db.query(Job).filter(Job.id == request.job_id).first()
        
        if not job:
            ctx_logger.warning(
                "Job not found",
                extra={"job_id": request.job_id}
            )
            
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Job with ID {request.job_id} not found"
            )
        
        ctx_logger.debug(
            "Job validated successfully",
            extra={"job_title": job.title}
        )
        
        # 2. Check for duplicate
        ctx_logger.debug("Checking for duplicate candidate")
        
        with PerformanceLogger(logger, "check_duplicate", request_id=request_id):
            existing = db.query(Candidate).filter(
                Candidate.email == request.email
            ).first()
        
        if existing:
            ctx_logger.warning(
                "Duplicate candidate found",
                extra={
                    "existing_candidate_id": existing.id,
                    "existing_name": existing.name
                }
            )
            
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=f"Candidate with email {request.email} already exists"
            )
        
        # 3. Create candidate
        candidate_data = {
            "name": request.name,
            "email": request.email,
            "phone": request.phone,
            "job_id": request.job_id,
            "resume_url": request.resume_url,
            "ats_candidate_id": request.ats_candidate_id,
            "source": request.source
        }
        
        # 4. If resume text provided, create with embeddings
        if request.resume_text:
            ctx_logger.info(
                "Creating candidate with resume embeddings",
                extra={
                    "resume_length": len(request.resume_text),
                    "resume_word_count": len(request.resume_text.split())
                }
            )
            
            with PerformanceLogger(logger, "create_candidate_with_embeddings", request_id=request_id):
                candidate = create_candidate_with_resume(
                    db,
                    request.name,
                    request.email,
                    request.resume_text,
                    request.job_id
                )
        else:
            ctx_logger.info("Creating candidate without resume")
            
            with PerformanceLogger(logger, "create_candidate", request_id=request_id):
                candidate = Candidate(**candidate_data)
                db.add(candidate)
                db.commit()
                db.refresh(candidate)
        
        ctx_logger.info(
            "Candidate created successfully",
            extra={
                "candidate_id": candidate.id,
                "has_embeddings": bool(request.resume_text)
            }
        )
        
        return candidate
    
    except HTTPException:
        # Re-raise HTTP exceptions without logging as errors
        raise
    
    except Exception as e:
        log_exception(logger, e, {
            "operation": "create_candidate",
            "request_id": request_id,
            "candidate_email": request.email,
            "job_id": request.job_id
        })
        
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to create candidate"
        )

# UPLOAD RESUME

@router.post("/{candidate_id}/resume", response_model=CandidateResponse)
async def upload_resume(
    candidate_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_session)
):
    """
    Upload and parse a resume file with comprehensive logging.
    
    Flow:
    1. Validate candidate exists
    2. Read file based on type (PDF or DOCX)
    3. Extract text
    4. Generate embeddings for RAG
    5. Update candidate record
    """
    
    request_id = f"req-{uuid.uuid4().hex[:12]}"
    
    ctx_logger = ContextLogger(
        logger=logger,
        request_id=request_id,
        candidate_id=candidate_id
    )
    
    ctx_logger.info(
        "Processing resume upload",
        extra={
            "filename": file.filename,
            "content_type": file.content_type,
            "file_size": file.size if hasattr(file, 'size') else None
        }
    )
    
    try:
        # 1. Find candidate
        ctx_logger.debug("Looking up candidate")
        
        with PerformanceLogger(logger, "lookup_candidate", request_id=request_id):
            candidate = db.query(Candidate).filter(
                Candidate.id == candidate_id
            ).first()
        
        if not candidate:
            ctx_logger.warning("Candidate not found")
            
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Candidate not found"
            )
        
        ctx_logger.info(
            "Candidate found",
            extra={
                "candidate_name": candidate.name,
                "candidate_email": candidate.email
            }
        )
        
        # 2. Validate file type
        allowed_types = [
            "application/pdf", 
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]
        
        if file.content_type not in allowed_types:
            ctx_logger.warning(
                "Invalid file type",
                extra={
                    "content_type": file.content_type,
                    "allowed_types": allowed_types
                }
            )
            
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Only PDF and DOCX files are supported"
            )
        
        # 3. Read file content
        ctx_logger.debug("Reading file content")
        
        with PerformanceLogger(logger, "read_file", request_id=request_id):
            content = await file.read()
        
        ctx_logger.info(
            "File read successfully",
            extra={"file_size_bytes": len(content)}
        )
        
        # 4. Extract text based on file type
        try:
            if file.content_type == "application/pdf":
                ctx_logger.debug("Extracting text from PDF")
                
                with PerformanceLogger(logger, "extract_pdf_text", request_id=request_id):
                    resume_text = extract_text_from_pdf(content)
            else:  # DOCX
                ctx_logger.debug("Extracting text from DOCX")
                
                with PerformanceLogger(logger, "extract_docx_text", request_id=request_id):
                    resume_text = extract_text_from_docx(content)
            
            ctx_logger.info(
                "Text extracted successfully",
                extra={
                    "text_length": len(resume_text),
                    "word_count": len(resume_text.split()),
                    "file_type": "PDF" if file.content_type == "application/pdf" else "DOCX"
                }
            )
        
        except Exception as e:
            log_exception(logger, e, {
                "operation": "extract_resume_text",
                "request_id": request_id,
                "candidate_id": candidate_id,
                "file_type": file.content_type
            })
            
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to parse resume: {str(e)}"
            )
        
        # 5. Validate extracted text
        if not resume_text or len(resume_text.strip()) < 50:
            ctx_logger.warning(
                "Resume text too short or empty",
                extra={"text_length": len(resume_text) if resume_text else 0}
            )
            
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Resume appears to be empty or too short"
            )
        
        # 6. Update candidate with resume text
        ctx_logger.debug("Updating candidate record with resume text")
        candidate.resume_text = resume_text
        
        # 7. Generate embeddings
        ctx_logger.info("Generating resume embeddings")
        
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        from database import ResumeEmbedding, get_embedding_function
        
        # Delete old embeddings
        ctx_logger.debug("Deleting old embeddings")
        
        with PerformanceLogger(logger, "delete_old_embeddings", request_id=request_id):
            deleted_count = db.query(ResumeEmbedding).filter(
                ResumeEmbedding.candidate_id == candidate_id
            ).delete()
        
        ctx_logger.debug(
            "Old embeddings deleted",
            extra={"deleted_count": deleted_count}
        )
        
        # Split into chunks
        ctx_logger.debug("Splitting resume into chunks")
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", ". ", " "]
        )
        
        with PerformanceLogger(logger, "chunk_resume", request_id=request_id):
            chunks = splitter.split_text(resume_text)
        
        ctx_logger.info(
            "Resume chunked",
            extra={
                "num_chunks": len(chunks),
                "avg_chunk_length": sum(len(c) for c in chunks) / len(chunks) if chunks else 0
            }
        )
        
        # Generate embeddings
        ctx_logger.info("Generating embeddings for chunks")
        
        embed_func = get_embedding_function()
        
        with PerformanceLogger(logger, "generate_embeddings", request_id=request_id):
            embeddings_list = embed_func.embed_documents(chunks)
        
        # Store embeddings
        ctx_logger.debug("Storing embeddings in database")
        
        with PerformanceLogger(logger, "store_embeddings", request_id=request_id):
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings_list)):
                resume_embedding = ResumeEmbedding(
                    candidate_id=candidate.id,
                    chunk_text=chunk,
                    chunk_metadata={"chunk_index": i},
                    embedding=embedding
                )
                db.add(resume_embedding)
            
            db.commit()
            db.refresh(candidate)
        
        ctx_logger.info(
            "Resume uploaded and processed successfully",
            extra={
                "candidate_id": candidate_id,
                "num_embeddings": len(chunks),
                "total_text_length": len(resume_text)
            }
        )
        
        return candidate
    
    except HTTPException:
        raise
    
    except Exception as e:
        log_exception(logger, e, {
            "operation": "upload_resume",
            "request_id": request_id,
            "candidate_id": candidate_id,
            "filename": file.filename
        })
        
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to process resume"
        )

# GET CANDIDATE

@router.get("/{candidate_id}", response_model=CandidateResponse)
async def get_candidate(
    candidate_id: int,
    db: Session = Depends(get_session)
):
    """
    Get candidate details by ID with logging.
    """
    
    request_id = f"req-{uuid.uuid4().hex[:12]}"
    
    logger.info(
        "Retrieving candidate details",
        extra={
            "request_id": request_id,
            "candidate_id": candidate_id
        }
    )
    
    try:
        with PerformanceLogger(logger, "get_candidate", request_id=request_id):
            candidate = db.query(Candidate).filter(
                Candidate.id == candidate_id
            ).first()
        
        if not candidate:
            logger.warning(
                "Candidate not found",
                extra={
                    "request_id": request_id,
                    "candidate_id": candidate_id
                }
            )
            
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Candidate not found"
            )
        
        logger.info(
            "Candidate retrieved successfully",
            extra={
                "request_id": request_id,
                "candidate_id": candidate_id,
                "candidate_name": candidate.name
            }
        )
        
        return candidate
    
    except HTTPException:
        raise
    
    except Exception as e:
        log_exception(logger, e, {
            "operation": "get_candidate",
            "request_id": request_id,
            "candidate_id": candidate_id
        })
        raise

@router.get("/", response_model=List[CandidateResponse])
async def list_candidates(
    job_id: Optional[int] = None,
    source: Optional[str] = None,
    limit: int = 50,
    offset: int = 0,
    db: Session = Depends(get_session)
):
    """
    List candidates with optional filters and logging.
    
    Filters:
    - job_id: Filter by job
    - source: Filter by source (manual, lever, greenhouse)
    - limit/offset: Pagination
    """
    
    request_id = f"req-{uuid.uuid4().hex[:12]}"
    
    logger.info(
        "Listing candidates",
        extra={
            "request_id": request_id,
            "job_id": job_id,
            "source": source,
            "limit": limit,
            "offset": offset
        }
    )
    
    try:
        with PerformanceLogger(logger, "list_candidates", request_id=request_id):
            query = db.query(Candidate)
            
            if job_id:
                query = query.filter(Candidate.job_id == job_id)
            
            if source:
                query = query.filter(Candidate.source == source)
            
            query = query.order_by(Candidate.created_at.desc())
            
            candidates = query.offset(offset).limit(limit).all()
        
        logger.info(
            "Candidates retrieved successfully",
            extra={
                "request_id": request_id,
                "count": len(candidates),
                "filters_applied": bool(job_id or source)
            }
        )
        
        return candidates
    
    except Exception as e:
        log_exception(logger, e, {
            "operation": "list_candidates",
            "request_id": request_id,
            "job_id": job_id,
            "source": source
        })
        raise

# HELPER FUNCTIONS WITH LOGGING

def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from PDF bytes with logging.
    
    Why: Resume parsing for RAG.
    Uses PyPDF2 for text extraction.
    """
    logger.debug("Starting PDF text extraction")
    
    try:
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        page_count = len(pdf_reader.pages)
        
        logger.debug(
            "PDF loaded successfully",
            extra={"page_count": page_count}
        )
        
        text = ""
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            text += page_text + "\n"
            
            if i == 0:
                logger.debug(
                    "First page extracted",
                    extra={"first_page_length": len(page_text)}
                )
        
        logger.debug(
            "PDF extraction completed",
            extra={
                "total_length": len(text),
                "pages_processed": page_count
            }
        )
        
        return text.strip()
    
    except Exception as e:
        log_exception(logger, e, {"operation": "extract_text_from_pdf"})
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(content: bytes) -> str:
    """
    Extract text from DOCX bytes with logging.
    
    Why: Resume parsing for RAG.
    Uses python-docx for text extraction.
    """
    logger.debug("Starting DOCX text extraction")
    
    try:
        docx_file = io.BytesIO(content)
        doc = docx.Document(docx_file)
        
        paragraph_count = len(doc.paragraphs)
        
        logger.debug(
            "DOCX loaded successfully",
            extra={"paragraph_count": paragraph_count}
        )
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        logger.debug(
            "DOCX extraction completed",
            extra={
                "total_length": len(text),
                "paragraphs_processed": paragraph_count
            }
        )
        
        return text.strip()
    
    except Exception as e:
        log_exception(logger, e, {"operation": "extract_text_from_docx"})
        raise ValueError(f"Failed to parse DOCX: {str(e)}")

# MODULE-LEVEL STATUS

def log_router_status():
    """Log the status of the candidates router."""
    logger.info(
        "Candidates API Router Status",
        extra={
            "prefix": "/api/candidates",
            "endpoints": [
                "POST /",
                "POST /{candidate_id}/resume",
                "GET /{candidate_id}",
                "GET /"
            ]
        }
    )


if __name__ == "__main__":
    logger.info("Candidates API module loaded")
    log_router_status()